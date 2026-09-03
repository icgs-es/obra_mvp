from __future__ import annotations

import csv
import io
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageOps
from pypdf import PdfReader


EXTRACTOR_VERSION = "intasa-doc-v1"
MAX_EXTRACTED_CHARS = 100_000
MAX_PDF_PAGES = 50
MAX_OCR_PAGES = 15
MAX_EXCEL_SHEETS = 10
MAX_EXCEL_ROWS = 10_000
MAX_EXCEL_COLUMNS = 50
MAX_EXCEL_CELLS = 50_000
SUBPROCESS_TIMEOUT = 90


class DocumentProcessingError(RuntimeError):
    def __init__(self, code):
        super().__init__(code)
        self.code = str(code)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    method: str
    page_count: int | None = None
    sheet_count: int | None = None
    ocr_used: bool = False
    summary: str = ""


def _limited(text):
    value = str(text or "").replace("\x00", "").strip()
    return value[:MAX_EXTRACTED_CHARS]


def _run(command):
    try:
        result = subprocess.run(
            command, capture_output=True, timeout=SUBPROCESS_TIMEOUT, check=False
        )
    except subprocess.TimeoutExpired as exc:
        raise DocumentProcessingError("extractor_timeout") from exc
    if result.returncode != 0:
        raise DocumentProcessingError("extractor_failed")
    return result.stdout.decode("utf-8", errors="replace")


def _ocr_image(path):
    return _run(["tesseract", str(path), "stdout", "-l", "spa+eng", "--psm", "6"])


def _extract_pdf(path):
    try:
        reader = PdfReader(str(path), strict=False)
        pages = len(reader.pages)
    except Exception as exc:
        raise DocumentProcessingError("pdf_corrupt") from exc
    if pages > MAX_PDF_PAGES:
        raise DocumentProcessingError("pdf_page_limit")
    direct_parts = []
    for page in reader.pages:
        try:
            direct_parts.append(page.extract_text() or "")
        except Exception:
            direct_parts.append("")
    direct = _limited("\n\n".join(direct_parts))
    meaningful = len("".join(direct.split())) >= max(20, pages * 15)
    if meaningful:
        return ExtractionResult(direct, "pypdf", page_count=pages)

    with tempfile.TemporaryDirectory(prefix="intasa_ia_ocr_") as temp_dir:
        prefix = Path(temp_dir) / "page"
        _run([
            "pdftoppm", "-f", "1", "-l", str(min(pages, MAX_OCR_PAGES)),
            "-jpeg", "-r", "200", str(path), str(prefix),
        ])
        parts = [_ocr_image(image) for image in sorted(Path(temp_dir).glob("page-*.jpg"))]
    text = _limited("\n\n".join(parts))
    if not text:
        raise DocumentProcessingError("ocr_empty")
    return ExtractionResult(
        text, "pypdf+tesseract", page_count=pages, ocr_used=True,
        summary=f"OCR aplicado a {min(pages, MAX_OCR_PAGES)} página(s).",
    )


def _extract_image(path):
    try:
        with Image.open(path) as source:
            source.verify()
        with Image.open(path) as source:
            image = ImageOps.exif_transpose(source).convert("RGB")
            with tempfile.TemporaryDirectory(prefix="intasa_ia_image_") as temp_dir:
                normalized = Path(temp_dir) / "normalized.png"
                image.save(normalized, format="PNG", optimize=True)
                text = _limited(_ocr_image(normalized))
    except (Image.DecompressionBombError, Image.DecompressionBombWarning) as exc:
        raise DocumentProcessingError("image_decompression_bomb") from exc
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError("image_corrupt") from exc
    if not text:
        raise DocumentProcessingError("ocr_empty")
    return ExtractionResult(text, "pillow+tesseract", page_count=1, ocr_used=True)


def _extract_docx(path):
    try:
        from docx import Document
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    except Exception as exc:
        raise DocumentProcessingError("docx_corrupt") from exc
    return ExtractionResult(_limited("\n".join(parts)), "python-docx")


def _safe_cell(value):
    text = str(value if value is not None else "").replace("\r", " ").replace("\n", " ")
    if text.lstrip().startswith(("=", "+", "-", "@")):
        text = "'" + text
    return text


def _extract_xlsx(path):
    try:
        from openpyxl import load_workbook
        workbook = load_workbook(str(path), read_only=True, data_only=True, keep_links=False)
        names = workbook.sheetnames[:MAX_EXCEL_SHEETS]
        parts, cells = [], 0
        for name in names:
            sheet = workbook[name]
            parts.append(f"[Hoja: {name}]")
            for row in sheet.iter_rows(max_row=MAX_EXCEL_ROWS, max_col=MAX_EXCEL_COLUMNS, values_only=True):
                cells += len(row)
                if cells > MAX_EXCEL_CELLS:
                    break
                parts.append(" | ".join(_safe_cell(value) for value in row))
            if cells > MAX_EXCEL_CELLS:
                break
        workbook.close()
    except Exception as exc:
        raise DocumentProcessingError("xlsx_corrupt") from exc
    return ExtractionResult(_limited("\n".join(parts)), "openpyxl", sheet_count=len(names))


def _extract_xls(path):
    try:
        import xlrd
        workbook = xlrd.open_workbook(str(path), on_demand=True)
        names = workbook.sheet_names()[:MAX_EXCEL_SHEETS]
        parts, cells = [], 0
        for name in names:
            sheet = workbook.sheet_by_name(name)
            parts.append(f"[Hoja: {name}]")
            for row_index in range(min(sheet.nrows, MAX_EXCEL_ROWS)):
                count = min(sheet.ncols, MAX_EXCEL_COLUMNS)
                cells += count
                if cells > MAX_EXCEL_CELLS:
                    break
                parts.append(" | ".join(_safe_cell(sheet.cell_value(row_index, col)) for col in range(count)))
            if cells > MAX_EXCEL_CELLS:
                break
        workbook.release_resources()
    except Exception as exc:
        raise DocumentProcessingError("xls_corrupt") from exc
    return ExtractionResult(_limited("\n".join(parts)), "xlrd", sheet_count=len(names))


def _decode_text(raw):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError("text_encoding")


def _extract_text(path, csv_mode=False):
    raw = Path(path).read_bytes()[:10_485_761]
    if b"\x00" in raw:
        raise DocumentProcessingError("binary_text")
    text, encoding = _decode_text(raw)
    if not csv_mode:
        return ExtractionResult(_limited(text), f"text:{encoding}")
    try:
        dialect = csv.Sniffer().sniff(text[:8192], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = []
    for index, row in enumerate(csv.reader(io.StringIO(text), dialect)):
        if index >= MAX_EXCEL_ROWS:
            break
        rows.append(" | ".join(_safe_cell(value) for value in row[:MAX_EXCEL_COLUMNS]))
    return ExtractionResult(_limited("\n".join(rows)), f"csv:{encoding}")


def extract_attachment(attachment):
    path = Path(attachment.file.path)
    extension = attachment.extension.lower()
    if extension == ".pdf":
        return _extract_pdf(path)
    if extension in {".jpg", ".jpeg", ".png", ".webp"}:
        return _extract_image(path)
    if extension == ".docx":
        return _extract_docx(path)
    if extension == ".xlsx":
        return _extract_xlsx(path)
    if extension == ".xls":
        return _extract_xls(path)
    if extension == ".csv":
        return _extract_text(path, csv_mode=True)
    if extension == ".txt":
        return _extract_text(path)
    raise DocumentProcessingError("unsupported_format")
