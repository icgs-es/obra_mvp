import csv
import io
import os
import shutil
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Permission
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas

from usuarios.models import Team

from .document_processing import extract_attachment
from .invoice_role_analysis import analyze_invoice_text
from .models import AdjuntoIA, ConversacionIA, MensajeIA, ProcesamientoMensajeIA
from .private_storage import private_ia_storage
from .tasks import process_document_message


User = get_user_model()


def _font(size=44):
    path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    return ImageFont.truetype(path, size) if os.path.exists(path) else ImageFont.load_default()


class ExtractionFormatTests(TestCase):
    def setUp(self):
        self.temp = tempfile.mkdtemp(prefix="ia_extract_")

    def tearDown(self):
        shutil.rmtree(self.temp, ignore_errors=True)

    def candidate(self, name, data):
        path = Path(self.temp) / name
        path.write_bytes(data)
        return SimpleNamespace(extension=path.suffix, file=SimpleNamespace(path=str(path)))

    def test_pdf_digital(self):
        stream = io.BytesIO()
        pdf = canvas.Canvas(stream)
        pdf.drawString(72, 750, "FACTURA F-2026-001 TOTAL 1210 EUR")
        pdf.save()
        result = extract_attachment(self.candidate("invoice.pdf", stream.getvalue()))
        self.assertIn("F-2026-001", result.text)
        self.assertFalse(result.ocr_used)

    def test_pdf_scanned_uses_ocr_path(self):
        image = Image.new("RGB", (1200, 400), "white")
        ImageDraw.Draw(image).text((40, 120), "FACTURA ESCANEADA 7788", fill="black", font=_font())
        stream = io.BytesIO()
        image.save(stream, format="PDF", resolution=150)
        with patch("intasa_ia.document_processing._ocr_image", return_value="FACTURA ESCANEADA 7788"):
            result = extract_attachment(self.candidate("scan.pdf", stream.getvalue()))
        self.assertTrue(result.ocr_used)
        self.assertIn("7788", result.text)

    def test_image_ocr(self):
        image = Image.new("RGB", (1400, 350), "white")
        ImageDraw.Draw(image).text((40, 100), "FACTURA IMAGEN 445566", fill="black", font=_font(54))
        stream = io.BytesIO()
        image.save(stream, format="PNG")
        result = extract_attachment(self.candidate("image.png", stream.getvalue()))
        self.assertTrue(result.ocr_used)
        self.assertIn("445566", result.text.replace(" ", ""))

    def test_docx_paragraphs_and_table(self):
        from docx import Document
        document = Document()
        document.add_paragraph("Proveedor Sintético")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Total"
        table.cell(0, 1).text = "121,00"
        stream = io.BytesIO()
        document.save(stream)
        result = extract_attachment(self.candidate("invoice.docx", stream.getvalue()))
        self.assertIn("Proveedor Sintético", result.text)
        self.assertIn("Total | 121,00", result.text)

    def test_xlsx_multiple_sheets_formula_not_evaluated(self):
        from openpyxl import Workbook
        workbook = Workbook()
        workbook.active.title = "Factura"
        workbook.active.append(["Número", "F-9"])
        workbook.create_sheet("Pagos").append(["Total", 99])
        workbook.active["C1"] = "=1+1"
        stream = io.BytesIO()
        workbook.save(stream)
        result = extract_attachment(self.candidate("book.xlsx", stream.getvalue()))
        self.assertEqual(result.sheet_count, 2)
        self.assertIn("[Hoja: Factura]", result.text)
        self.assertNotIn("=1+1", result.text)

    def test_csv_and_txt_neutralize_formulas(self):
        csv_result = extract_attachment(self.candidate("table.csv", b"name,value\nrow,=CMD()\n"))
        text_result = extract_attachment(self.candidate("note.txt", "Texto seguro ñ".encode()))
        self.assertIn("'=CMD()", csv_result.text)
        self.assertIn("Texto seguro", text_result.text)

    def test_invoice_roles_never_promote_customer_to_supplier(self):
        text = """FACTURA\nVENDEDOR\n1252/2026\n3314\nADRI MARTIN INVESTMENT SL\nCALLE LOS HERREROS DE SOSTOA 66\nCLIENTE\n27/08/2026\nB93578649\nESCAYOLAS ZAMORA S.L.\nB49155609\nescayolaszamora@escayolaszamora.es\nTOTAL FACTURA\n214,85\n21,00 45,12\n259,97 EUR\ntransferencia caja rural ES42"""
        result = analyze_invoice_text(text)
        self.assertNotEqual(result["supplier_name"], "ADRI MARTIN INVESTMENT SL")
        self.assertEqual(result["customer_name"], "ADRI MARTIN INVESTMENT SL")
        self.assertEqual(result["invoice_number"], "1252/2026")
        self.assertEqual(result["invoice_date"], "27/08/2026")
        self.assertEqual(result["taxable_base"], 214.85)
        self.assertEqual(result["vat_amount"], 45.12)
        self.assertEqual(result["total"], 259.97)

    def test_unknown_supplier_fallback_preserves_customer(self):
        result = analyze_invoice_text(
            "FACTURA\n1252/2026\nCLIENTE\nADRI MARTIN INVESTMENT SL\nTOTAL 259,97 EUR"
        )
        self.assertIsNone(result["supplier_name"])
        self.assertEqual(result["customer_name"], "ADRI MARTIN INVESTMENT SL")
        self.assertIn("Proveedor/emisor no identificado", result["warnings"][0])


@override_settings(SECURE_SSL_REDIRECT=False)
class ProcessingWorkflowTests(TestCase):
    def setUp(self):
        self.temp = tempfile.mkdtemp(prefix="ia_process_")
        private_ia_storage._location = self.temp
        for cached in ("base_location", "location", "base_url"):
            private_ia_storage.__dict__.pop(cached, None)
        self.team = Team.objects.create(name="PROCESSING TEST")
        self.owner = User.objects.create_user(username="processing-owner")
        self.shared = User.objects.create_user(username="processing-shared")
        permission = Permission.objects.get(content_type__app_label="intasa_ia", codename="use_intasa_ia")
        self.owner.user_permissions.add(permission)
        self.shared.user_permissions.add(permission)
        self.conversation = ConversacionIA.objects.create(user=self.owner, titulo="Factura")
        self.message = MensajeIA.objects.create(
            conversacion=self.conversation, rol=MensajeIA.Rol.USUARIO,
            contenido="¿Me analizas esta factura?",
        )
        self.attachment = AdjuntoIA.objects.create(
            conversation=self.conversation, message=self.message, owner=self.owner,
            file=SimpleUploadedFile("invoice.txt", b"FACTURA F-123 TOTAL 100 EUR", content_type="text/plain"),
            original_name="invoice.txt", safe_display_name="invoice.txt",
            declared_mime="text/plain", detected_mime="text/plain", extension=".txt",
            size_bytes=27, sha256="c" * 64, status=AdjuntoIA.Estado.UPLOADED,
        )
        self.processing = ProcesamientoMensajeIA.objects.create(message=self.message)

    def tearDown(self):
        shutil.rmtree(self.temp, ignore_errors=True)

    @patch("intasa_ia.tasks.generar_respuesta_segura")
    def test_task_extracts_then_calls_text_responses_once_without_web_or_files(self, generate):
        generate.return_value = {
            "contenido": "La factura invoice.txt es F-123 por 100 EUR.",
            "proveedor": "mock", "modelo": "mock-model", "request_id": "mock-1",
            "tokens_entrada": 10, "tokens_salida": 5,
            "metadata": {"external_call": False, "store": False},
        }
        process_document_message.run(self.processing.pk, str(self.processing.task_key))
        self.processing.refresh_from_db()
        self.attachment.refresh_from_db()
        self.assertEqual(self.processing.status, ProcesamientoMensajeIA.Estado.COMPLETED)
        self.assertEqual(self.attachment.status, AdjuntoIA.Estado.READY)
        self.assertIn("F-123", self.attachment.extracted_text)
        generate.assert_called_once()
        kwargs = generate.call_args.kwargs
        self.assertFalse(kwargs["web_search_enabled"])
        self.assertIn("DOCUMENTO_NO_CONFIABLE", kwargs["document_context"])
        self.assertNotIn("input_file", repr(kwargs))
        self.assertNotIn("input_image", repr(kwargs))
        self.assertIn("invoice.txt", self.processing.assistant_message.contenido)

    @patch("intasa_ia.tasks.generar_respuesta_segura")
    def test_duplicate_task_does_not_duplicate_provider_or_answer(self, generate):
        generate.return_value = {
            "contenido": "OK", "proveedor": "mock", "modelo": "mock", "request_id": "1",
            "tokens_entrada": 1, "tokens_salida": 1, "metadata": {"store": False},
        }
        process_document_message.run(self.processing.pk, str(self.processing.task_key))
        process_document_message.run(self.processing.pk, str(self.processing.task_key))
        self.assertEqual(generate.call_count, 1)
        self.assertEqual(MensajeIA.objects.filter(rol=MensajeIA.Rol.ASISTENTE).count(), 1)

    def test_polling_owner_and_unauthorized_scope(self):
        self.client.force_login(self.owner)
        url = reverse("intasa_ia:estado_procesamiento", args=[self.conversation.pk, self.message.pk])
        self.assertEqual(self.client.get(url).status_code, 200)
        self.client.force_login(self.shared)
        self.assertEqual(self.client.get(url).status_code, 404)

    def test_conversation_delete_purges_extracted_text_record(self):
        self.attachment.extracted_text = "sensitive synthetic content"
        self.attachment.save(update_fields=("extracted_text",))
        self.conversation.delete()
        self.assertFalse(AdjuntoIA.objects.filter(pk=self.attachment.pk).exists())
        self.assertFalse(ProcesamientoMensajeIA.objects.filter(pk=self.processing.pk).exists())
